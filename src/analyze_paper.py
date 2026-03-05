
import docx
import sys

def analyze_paper_structure(file_path):
    doc = docx.Document(file_path)
    content = []
    
    # 概览段落
    content.append(f"Total Paragraphs: {len(doc.paragraphs)}")
    content.append("--- Paragraph Samples ---")
    for i, para in enumerate(doc.paragraphs[:20]): # 只看前20个段落
        if para.text.strip():
            content.append(f"P{i} (Style: {para.style.name}): {para.text[:100]}...")
            
    # 概览表格
    content.append(f"\nTotal Tables: {len(doc.tables)}")
    for i, table in enumerate(doc.tables):
        content.append(f"Table {i}: {len(table.rows)} rows, {len(table.columns)} cols")
            
    return '\n'.join(content)

if __name__ == "__main__":
    file_path = r"C:\Users\Administrator\Desktop\小论文2\FEA\paper_final_v2_tablegrid.docx"
    try:
        sys.stdout.reconfigure(encoding='utf-8')
        print(analyze_paper_structure(file_path))
    except Exception as e:
        print(f"Error: {e}")
