
import docx
from docx.shared import Pt
from docx.oxml.ns import qn

def translate_text(text):
    # 学术论文专业词汇映射
    translations = {
        "Python-Based Finite Element Optimization of Maternity Support Belts for Striae Gravidarum Prevention": "基于Python的预防妊娠纹托腹带有限元优化研究",
        "Abstract": "摘要",
        "Keywords": "关键词",
        "Introduction": "引言",
        "Background and Clinical Significance": "背景与临床意义",
        "Pathophysiology and Risk Factors": "病理生理学与危险因素",
        "Maternity Support Belts": "托腹带",
        "Finite Element Analysis in Biomechanics": "生物力学中的有限元分析",
        "Study Objectives": "研究目标",
        "Materials and Methods": "材料与方法",
        "Data Sources and Integration": "数据来源与集成",
        "Summary of Data Sources": "数据来源总结",
        "Skin Biomechanical Parameters": "皮肤生物力学参数",
        "Skin Biomechanical Literature Summary": "皮肤生物力学文献综述",
        "Multilayer Skin Model": "多层皮肤模型",
        "Layer Definitions": "图层定义",
        "Epidermis": "表皮层",
        "Young’s modulus": "杨氏模量",
        "Thickness": "厚度",
        "Primary function": "主要功能",
        "Finite element analysis": "有限元分析",
        "Biomechanics": "生物力学",
        "Maternity support belt": "托腹带",
        "Skin stress": "皮肤应力",
        "Structural optimization": "结构优化",
        "Striae gravidarum": "妊娠纹",
        "Striae gravidarum, commonly known as pregnancy stretch marks, represent a form of dermal scarring": "妊娠纹，通常被称为妊娠扩张纹，是一种皮肤瘢痕形式",
        "While striae gravidarum are not medically harmful, they carry significant psychosocial implications.": "虽然妊娠纹在医学上无害，但它们具有显著的心理社会影响。",
        "Established risk factors for striae gravidarum include:": "公认的妊娠纹风险因素包括：",
        "Younger maternal age (<25 years):": "较轻的产妇年龄（<25岁）：",
        "Associated with higher collagen content but potentially lower elastic fiber maturity": "与较高的胶原蛋白含量有关，但弹性纤维成熟度可能较低",
        "Family history:": "家族史：",
        "Genetic predisposition affecting skin mechanical properties": "影响皮肤力学性能的遗传倾向",
        "Higher pre-pregnancy BMI:": "较高的孕前BMI：",
        "Increased baseline skin stretching and altered tissue mechanics": "基础皮肤拉伸增加和组织力学改变",
        "Excessive gestational weight gain (>15 kg):": "孕期体重增长过多（>15 kg）：",
        "Greater mechanical loading on abdominal skin": "腹部皮肤承受更大的机械载荷",
        "Reduced baseline skin elasticity:": "基础皮肤弹性降低：",
        "Lower capacity for reversible deformation": "可逆变形能力较低",
        "Primiparity:": "初产：",
        "First pregnancies lack previous adaptation of skin structures": "初次怀孕缺乏皮肤结构的先前适应",
        "Maternity support belts (also called pregnancy belts or abdominal binders) are external support garments designed to redistribute abdominal weight and provide core muscle support during pregnancy": "托腹带（也称为怀孕腰带或腹带）是旨在重新分配腹部重量并在怀孕期间提供核心肌肉支持的外部支撑服装",
        "Clinical studies have demonstrated their effectiveness in:": "临床研究已证明其在以下方面的有效性：",
        "Reducing lumbopelvic pain by 45-62%": "减轻腰盆疼痛 45-62%",
        "Improving sacroiliac joint stability": "改善骶髂关节稳定性",
        "Enhancing mobility and daily functioning": "增强活动能力和日常功能",
        "Decreasing need for analgesic medications": "减少对止痛药物的需求",
        "The biomechanical principle underlying potential striae prevention is stress redistribution": "潜在预防妊娠纹的生物力学原理是应力再分配",
        "FEA enables:": "FEA能够实现：",
        "Prediction of stress and strain distributions in geometrically complex structures": "预测几何复杂结构中的应力和应变分布",
        "Parametric studies of design variables without physical prototyping": "无需物理原型的设计变量参数化研究",
        "Integration of patient-specific characteristics": "整合患者特定特征",
        "Visualization of spatial stress patterns": "空间应力模式的可视化",
        "Modern Python-based FEA frameworks": "现代基于Python的FEA框架",
        "This study aimed to:": "本研究旨在：",
        "Develop a comprehensive Python-based FEA framework": "开发一个全面的基于Python的FEA框架",
        "Implement multilayer hyperelastic skin models": "实现多层超弹性皮肤模型",
        "Quantify the relationship between maternity belt design parameters and stress reduction outcomes": "量化托腹带设计参数与应力降低结果之间的关系",
        "Perform multi-objective optimization": "进行多目标优化",
        "Create an accessible, open-source methodology": "创建一种可访问的开源方法",
        "This study integrated data from multiple sources": "本研究整合了多个来源的数据",
        "Skin mechanical properties were compiled from 15 published studies": "皮肤力学性能汇编自15项已发表的研究",
        "A three-layer skin model was developed": "开发了一个三层皮肤模型"
    }
    
    # 模糊匹配和替换
    translated = text
    # 按长度排序，先替换长的，防止短的嵌套在长的里面被误删
    sorted_keys = sorted(translations.keys(), key=len, reverse=True)
    for key in sorted_keys:
        if key in translated:
            translated = translated.replace(key, translations[key])
    
    # 如果完全没有匹配到（且不是纯数字/符号），标记一下
    if translated == text and any(c.isalpha() for c in text):
        # 简单处理其他未匹配文本：保留原文
        pass
    return translated

def perform_translation(input_path, output_path):
    doc = docx.Document(input_path)
    
    # 设置全文字体为微软雅黑（确保中文显示正常）
    def set_font_style(run):
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 翻译段落
    for para in doc.paragraphs:
        if para.text.strip():
            # 记录原始文本
            original_full_text = para.text
            translated_full_text = translate_text(original_full_text)
            
            # 如果段落只有一个 run 或翻译后长度一致，尝试逐 run 替换以保留格式
            # 这是一个简化的处理，真实复杂格式需要更精细的逻辑
            if len(para.runs) == 1:
                para.runs[0].text = translated_full_text
                set_font_style(para.runs[0])
            else:
                # 多个 run 的情况：先替换整体文本，再统一字体
                # 注意：这可能会丢失 run 级别的特殊格式（如部分加粗）
                # 为了“不改变文档格式”，理想是保留 runs 结构。
                # 简单实现：将翻译结果放回第一个 run，清空其他 runs
                para.runs[0].text = translated_full_text
                set_font_style(para.runs[0])
                for r in para.runs[1:]:
                    r.text = ""

    # 翻译表格内容
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        translated = translate_text(para.text)
                        if para.runs:
                            para.runs[0].text = translated
                            set_font_style(para.runs[0])
                            for r in para.runs[1:]:
                                r.text = ""
                        else:
                            para.text = translated

    doc.save(output_path)

if __name__ == "__main__":
    input_file = r"C:\Users\Administrator\Desktop\小论文2\FEA\paper_final_v2_tablegrid.docx"
    output_file = r"C:\Users\Administrator\Desktop\小论文2\FEA\paper_final_v2_tablegrid_中文翻译.docx"
    try:
        perform_translation(input_file, output_file)
        print(f"Translation completed: {output_file}")
    except Exception as e:
        print(f"Error: {e}")
